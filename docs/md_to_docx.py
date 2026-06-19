"""
Convert screens_specification.md -> screens_specification.docx
Styled: headings, bold, italic, bullet lists, horizontal rules.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_FILE = "screens_specification.md"
OUT_FILE = "screens_specification.docx"

# ── Colour palette ─────────────────────────────────────────────────────────────
C_H1   = RGBColor(0x0D, 0x47, 0xA1)   # deep blue
C_H2   = RGBColor(0x15, 0x65, 0xC0)   # medium blue
C_H3   = RGBColor(0x19, 0x76, 0xD2)   # lighter blue
C_BODY = RGBColor(0x21, 0x21, 0x21)   # near-black
C_RULE = RGBColor(0xBD, 0xBD, 0xBD)   # light grey for divider

FONT_NAME = "Calibri"


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_font(run, size_pt, bold=False, italic=False, color=None, name=None):
    run.font.name     = name or FONT_NAME
    run.font.size     = Pt(size_pt)
    run.font.bold     = bold
    run.font.italic   = italic
    if color:
        run.font.color.rgb = color
    # East-Asian font slot (Vietnamese glyphs)
    rPr = run._r.get_or_add_rPr()
    eastAsian = OxmlElement("w:rFonts")
    eastAsian.set(qn("w:eastAsia"), name or FONT_NAME)
    rPr.insert(0, eastAsian)


def add_paragraph_spacing(para, before=0, after=0, line_spacing=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line_spacing:
        pf.line_spacing = Pt(line_spacing)


def add_horizontal_rule(doc):
    """Insert a thin grey horizontal line."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BDBDBD")
    pBdr.append(bottom)
    pPr.append(pBdr)


def inline_style(para, text, base_size=11, base_color=None):
    """
    Parse inline markdown: **bold**, *italic*, `code`.
    Returns nothing – appends runs directly to para.
    """
    # Pattern captures: **…**, *…*, `…`, or plain text
    pattern = re.compile(r'\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+)', re.DOTALL)
    for m in pattern.finditer(text):
        bold_txt, it_txt, code_txt, plain_txt = m.groups()
        run = para.add_run()
        if bold_txt is not None:
            run.text = bold_txt
            set_font(run, base_size, bold=True, color=base_color)
        elif it_txt is not None:
            run.text = it_txt
            set_font(run, base_size, italic=True, color=base_color)
        elif code_txt is not None:
            run.text = code_txt
            set_font(run, base_size - 0.5, name="Courier New", color=RGBColor(0xC6, 0x28, 0x28))
        elif plain_txt is not None:
            run.text = plain_txt
            set_font(run, base_size, color=base_color)


# ── Main conversion ────────────────────────────────────────────────────────────

def convert(md_path, out_path):
    with open(md_path, encoding="utf-8") as f:
        lines = f.readlines()

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Default body style ──
    style = doc.styles["Normal"]
    style.font.name = FONT_NAME
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n").rstrip("\r")

        # ── Blank line ──
        if line.strip() == "":
            i += 1
            continue

        # ── Horizontal rule ---
        if re.match(r'^-{3,}$', line.strip()):
            add_horizontal_rule(doc)
            i += 1
            continue

        # ── H1  # Title
        if line.startswith("# ") and not line.startswith("## "):
            text = line[2:].strip()
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_paragraph_spacing(para, before=0, after=10, line_spacing=16)
            run = para.add_run(text)
            set_font(run, 18, bold=True, color=C_H1)
            i += 1
            continue

        # ── H3  ### Section
        if line.startswith("### "):
            text = line[4:].strip()
            para = doc.add_paragraph()
            add_paragraph_spacing(para, before=10, after=4, line_spacing=14)
            run = para.add_run(text)
            set_font(run, 12, bold=True, color=C_H3)
            i += 1
            continue

        # ── H2  ## Screen
        if line.startswith("## "):
            text = line[3:].strip()
            para = doc.add_paragraph()
            add_paragraph_spacing(para, before=14, after=4, line_spacing=16)
            run = para.add_run(text)
            set_font(run, 14, bold=True, color=C_H2)
            i += 1
            continue

        # ── Ordered list  1. …
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            number = m.group(1)
            content = m.group(2).strip()
            para = doc.add_paragraph(style="List Number")
            add_paragraph_spacing(para, before=2, after=2, line_spacing=13)
            para.paragraph_format.left_indent = Cm(0.5)
            inline_style(para, content, base_size=11, base_color=C_BODY)
            i += 1
            continue

        # ── Unordered sub-bullet   - …  or   * …
        m = re.match(r'^(\s+)[-*]\s+(.*)', line)
        if m:
            indent_depth = len(m.group(1))
            content = m.group(2).strip()
            para = doc.add_paragraph(style="List Bullet")
            add_paragraph_spacing(para, before=1, after=1, line_spacing=13)
            # indent proportional to spaces
            para.paragraph_format.left_indent = Cm(0.5 + indent_depth * 0.1)
            inline_style(para, content, base_size=10.5, base_color=C_BODY)
            i += 1
            continue

        # ── Top-level bullet  * …  (no indent)
        m = re.match(r'^[-*]\s+(.*)', line)
        if m:
            content = m.group(1).strip()
            para = doc.add_paragraph(style="List Bullet")
            add_paragraph_spacing(para, before=2, after=2, line_spacing=13)
            para.paragraph_format.left_indent = Cm(0.3)
            inline_style(para, content, base_size=11, base_color=C_BODY)
            i += 1
            continue

        # ── Plain paragraph ──
        para = doc.add_paragraph()
        add_paragraph_spacing(para, before=2, after=4, line_spacing=14)
        inline_style(para, line, base_size=11, base_color=C_BODY)
        i += 1

    doc.save(out_path)
    print(f"[OK] Saved: {out_path}")


if __name__ == "__main__":
    convert(MD_FILE, OUT_FILE)
